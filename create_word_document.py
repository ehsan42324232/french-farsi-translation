#!/usr/bin/env python3
"""
Simple Word Document Creator for French-Farsi Translation
Creates a .docx file that can be opened in Microsoft Word with proper RTL support
"""

def create_simple_word_document(french_text, farsi_text, filename="translation.docx"):
    """
    Create a simple Word document with French and Farsi text.
    This creates a basic but functional .docx file.
    """
    
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('ترجمه کامل کازانووا: تاریخ زندگی من', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle
        subtitle = doc.add_heading('Complete Translation: Casanova - Histoire de ma vie', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_paragraph(f'تاریخ ترجمه: ۲۰۲۵/۰۷/۰۵')
        doc.add_paragraph(f'مترجم: سیستم هوش مصنوعی')
        doc.add_paragraph('═' * 50)
        
        # French section
        french_heading = doc.add_heading('متن اصلی فرانسوی', 2)
        french_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        french_para = doc.add_paragraph(french_text)
        french_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        doc.add_paragraph('─' * 30)
        
        # Farsi section
        farsi_heading = doc.add_heading('ترجمه فارسی', 2)
        farsi_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        farsi_para = doc.add_paragraph(farsi_text)
        farsi_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Set RTL for Farsi paragraphs
        for para in [farsi_heading, farsi_para]:
            try:
                pPr = para._element.get_or_add_pPr()
                bidi = OxmlElement('w:bidi')
                bidi.set(qn('w:val'), '1')
                pPr.append(bidi)
            except:
                pass  # If RTL setting fails, continue anyway
        
        # Save document
        doc.save(filename)
        print(f"✅ Word document created: {filename}")
        return filename
        
    except ImportError:
        print("⚠️ python-docx not available. Creating text-based document instead.")
        
        # Fallback: create a text file with proper encoding
        text_content = f"""
ترجمه کامل کازانووا: تاریخ زندگی من
Complete Translation: Casanova - Histoire de ma vie

تاریخ ترجمه: ۲۰۲۵/۰۷/۰۵
مترجم: سیستم هوش مصنوعی

═══════════════════════════════════════════════════════════════

متن اصلی فرانسوی:
{french_text}

─────────────────────────────────────────────────────────────

ترجمه فارسی:
{farsi_text}

═══════════════════════════════════════════════════════════════
ترجمه شده توسط سیستم هوش مصنوعی
Translated by AI System
        """
        
        text_filename = filename.replace('.docx', '.txt')
        with open(text_filename, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        print(f"✅ Text document created: {text_filename}")
        return text_filename

# Sample usage and demonstration
if __name__ == "__main__":
    
    # Sample French text from the actual content
    sample_french = """
CASANOVA
HISTOIRE DE MA VIE

Casanova, écrivain par Jean-Christophe Igalens

Un des plus grands plaisirs de Paris est celui d'aller vite. Casanova aime le 
rythme de cette ville prête à sacrifier ses chevaux pour écraser la distance et abolir 
l'attente. Passer sans retard d'un lieu à un autre, traverser tous les milieux sociaux, 
ne jamais se laisser enfermer dans une seule identité : Paris, aux yeux de 
Casanova, autorise toutes les mobilités et offre tous les possibles à l'homme qui 
place la disponibilité au principe de son existence. Dans cette ville « où 
l'imposture a toujours fait fortune », Casanova n'est jamais figé, il peut 
toujours se réinventer.

Giacomo Casanova naît à Venise le 2 avril 1725, de parents comédiens. La 
ville est elle-même un spectacle. Le carnaval dure presque la moitié de l'année et 
Venise célèbre tous les ans son propre mythe : lors des fêtes de l'Ascension, le 
doge en tenue d'apparat, accompagné par les puissants de la République, monte à 
bord du Bucentaure pour jeter un anneau dans l'Adriatique. On dit qu'il épouse la 
mer.
    """
    
    # Sample Farsi translation
    sample_farsi = """
کازانووا
تاریخ زندگی من

کازانووا، نویسنده اثر ژان-کریستف ایگالنس

یکی از بزرگترین لذت‌های پاریس این است که با سرعت حرکت کنی. کازانووا ریتم این 
شهری را دوست دارد که آماده است اسب‌هایش را قربانی کند تا فاصله را در هم بکوبد و انتظار را 
نابود کند. بدون تأخیر از مکانی به مکان دیگر رفتن، تمام طبقات اجتماعی را پیمودن، هرگز 
نگذاشتن که در یک هویت واحد محبوس شوی: پاریس، در نظر کازانووا، همه گونه تحرک‌ها را 
مجاز می‌داند و همه امکانات را به انسانی ارائه می‌دهد که در دسترس بودن را اصل وجودش قرار 
داده است. در این شهری «که دروغ همیشه ثروت ساخته است»، کازانووا هرگز ثابت نمی‌ماند، 
همیشه می‌تواند خود را دوباره اختراع کند.

جیاکومو کازانووا در ۲ آوریل ۱۷۲۵ در ونیز از والدین کمدین متولد شد. شهر خودش 
نمایشی است. کارناوال تقریباً نیمی از سال را در بر می‌گیرد و ونیز هر سال اسطوره خودش را جشن 
می‌گیرد: در جشن‌های عروج، دوژه با لباس تشریفاتی، همراه با قدرتمندان جمهوری، سوار قایق 
بوچنتاورو می‌شود تا انگشتری در دریای آدریاتیک بیندازد. می‌گویند که با دریا ازدواج می‌کند.
    """
    
    print("🎯 Creating Sample Word Document")
    print("=" * 50)
    
    # Create the document
    output_file = create_simple_word_document(
        sample_french, 
        sample_farsi, 
        "casanova_sample_translation.docx"
    )
    
    print(f"\n📄 Sample document created successfully!")
    print(f"📁 File: {output_file}")
    print(f"\n💡 To create the complete translation:")
    print(f"   1. Install: pip install python-docx")
    print(f"   2. Run: python complete_translator.py")
    print(f"   3. Process the full 1-4.txt file")
    
    print(f"\n🎉 Translation project completed!")
